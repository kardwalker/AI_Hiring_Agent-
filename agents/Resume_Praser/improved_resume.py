from typing import List, Optional, Dict, Any
import os
import asyncio
import logging
import json
import httpx
from datetime import datetime, timedelta
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.vectorstores import FAISS
import re
import numpy as np
import argparse
import sys

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# PyMuPDF import with better error handling
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
    logger.info("✓ PyMuPDF available for hyperlink extraction")
except ImportError:
    FITZ_AVAILABLE = False
    logger.warning("⚠️  PyMuPDF not available. Install with: pip install PyMuPDF")

# Load environment variables
load_dotenv()

# Azure OpenAI configuration
AZURE_OPENAI_API_KEY = os.getenv("AZURE_API_KEY") 
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_ENDPOINT") 
AZURE_EMBEDDING_DEPLOYMENT = os.getenv("AZURE_EMBEDDING_DEPLOYMENT", "text-embedding-ada-002")

if not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT:
    logger.error("❌ Azure OpenAI credentials missing. Set AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in .env")

# Initialize models
try:
    chat_model = AzureChatOpenAI(
        api_key=AZURE_OPENAI_API_KEY,
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        deployment_name="gpt-4o-mini",
        api_version="2024-05-01-preview",
        temperature=0.2,
        max_tokens=512
    )
    embedding_model = AzureOpenAIEmbeddings(
    azure_endpoint="https://agentic-keran-framework.openai.azure.com/",
    api_key=os.getenv("Embedding_api_key")                    ,
    deployment="text-embedding-ada-002",  # Deployment name must match what you
    # created in Azure
    model="text-embedding-ada-002",       # Model name (optional, but helps for clarity)
    api_version="2023-05-15",        # Use correct version (based on Azure docs or trial-and-error))
)

   
    logger.info("✓ Azure OpenAI models initialized successfully")
except Exception as e:
    logger.error(f"❌ Failed to initialize Azure OpenAI models: {e}")
    chat_model = None
    embedding_model = None


def extract_hyperlinks_from_pdf(file_path: str) -> Dict[str, List[Dict[str, Any]]]:
    """Extract hyperlinks from PDF using PyMuPDF with improved error handling.
    
    Returns:
        Dict with categorized links and their metadata
    """
    default_result = {
        'github': [],
        'linkedin': [],
        'research_publications': [],
        'certifications': [],
        'social_media': [],
        'other': []
    }
    
    if not FITZ_AVAILABLE:
        logger.warning("PyMuPDF not available. Cannot extract hyperlinks from PDF.")
        return default_result
    
    if not file_path.lower().endswith('.pdf'):
        logger.info("File is not a PDF. Skipping hyperlink extraction.")
        return default_result
    
    if not os.path.exists(file_path):
        logger.error(f"PDF file not found: {file_path}")
        return default_result
    
    links = default_result.copy()
    
    try:
        logger.info(f"Opening PDF for hyperlink extraction: {file_path}")
        doc = fitz.open(file_path)
        total_links_found = 0
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            link_list = page.get_links()
            
            for link in link_list:
                if 'uri' in link and link['uri']:
                    url = link['uri']
                    categorized_url = categorize_link(url)
                    category = categorized_url['category']
                    
                    if category in links:
                        link_info = {
                            'url': url,
                            'page': page_num + 1,
                            'description': categorized_url['description'],
                            'rect': link.get('from', None)
                        }
                        links[category].append(link_info)
                        total_links_found += 1
        
        doc.close()
        
        # Remove duplicates while preserving order
        for category in links:
            seen = set()
            unique_links = []
            for link in links[category]:
                url = link['url']
                if url not in seen:
                    seen.add(url)
                    unique_links.append(link)
            links[category] = unique_links
        
        logger.info(f"✓ Extracted {total_links_found} total hyperlinks from PDF")
        for category, link_list in links.items():
            if link_list:
                logger.info(f"  - {category.replace('_', ' ').title()}: {len(link_list)} links")
        
        return links
        
    except Exception as e:
        logger.error(f"❌ Error extracting hyperlinks from PDF: {e}")
        return default_result


def categorize_link(url: str) -> Dict[str, str]:
    """Categorize a URL into relevant resume sections with improved patterns."""
    url_lower = url.lower()
    
    # GitHub links
    if 'github.com' in url_lower:
        if '/repositories' in url_lower or '/repos' in url_lower:
            return {'category': 'github', 'description': 'GitHub repositories'}
        else:
            return {'category': 'github', 'description': 'GitHub profile/repository'}
    
    # LinkedIn links
    elif 'linkedin.com' in url_lower:
        return {'category': 'linkedin', 'description': 'LinkedIn profile'}
    
    # Research and publication platforms
    elif any(platform in url_lower for platform in [
        'researchgate.net', 'scholar.google.', 'orcid.org', 'ieee.org',
        'acm.org', 'arxiv.org', 'pubmed.ncbi.nlm.nih.gov', 'semanticscholar.org',
        'dblp.org', 'springer.com', 'sciencedirect.com', 'jstor.org',
        'doi.org', 'researchgate.net'
    ]):
        return {'category': 'research_publications', 'description': 'Research/Publication platform'}
    
    # Certification platforms
    elif any(cert_platform in url_lower for cert_platform in [
        'credly.com', 'badgr.com', 'certifications.aws', 'cloud.google.com/certification',
        'docs.microsoft.com/certifications', 'coursera.org/account/accomplishments',
        'udacity.com/certificate', 'edx.org/certificates', 'codecademy.com/profiles',
        'freecodecamp.org/certification', 'cisco.com/c/en/us/training-events/training-certifications'
    ]):
        return {'category': 'certifications', 'description': 'Certification platform'}
    
    # Social media platforms
    elif any(social in url_lower for social in [
        'twitter.com', 'x.com', 'facebook.com', 'instagram.com', 'youtube.com',
        'medium.com', 'dev.to', 'stackoverflow.com', 'reddit.com'
    ]):
        return {'category': 'social_media', 'description': 'Social media profile'}
    
    # Portfolio/personal websites
    elif any(indicator in url_lower for indicator in [
        'portfolio', 'blog', 'personal', 'website'
    ]) or url_lower.count('.') == 1:
        return {'category': 'other', 'description': 'Personal website/portfolio'}
    
    else:
        return {'category': 'other', 'description': 'External link'}


def resume_loader(file_path: str) -> List[Document]:
    """Load resume from various file formats with improved error handling."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Loading {ext} file: {file_path}")
    
    try:
        if ext == ".pdf":
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            logger.info(f"✓ Loaded PDF with {len(docs)} pages")
            return docs
            
        elif ext in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            logger.info(f"✓ Loaded text file ({len(content)} characters)")
            return [Document(page_content=content, metadata={"source": file_path})]
            
        elif ext == ".docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                logger.info(f"✓ Loaded DOCX file ({len(text)} characters)")
                return [Document(page_content=text, metadata={"source": file_path})]
            except ImportError:
                raise ImportError("python-docx package required for .docx files. Install with: pip install python-docx")
        else:
            raise ValueError(f"Unsupported file format: {ext}. Supported: .pdf, .txt, .md, .docx")
            
    except Exception as e:
        logger.error(f"❌ Error loading file {file_path}: {e}")
        raise


class ResumeProcessor:
    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=50,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        
    def extract_contact_info(self, text: str, hyperlinks: Optional[Dict] = None) -> Dict[str, Any]:
        """Extract and return structured contact information with improved patterns."""
        contact_info = {}
        
        # Email pattern (improved)
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['emails'] = list(set(emails))  # Remove duplicates
        
        # Phone pattern (various formats, improved)
        phone_patterns = [
            r'(?:\+?91[-.\s]?)?(?:\(?([0-9]{3,4})\)?[-.\s]?)?([0-9]{3})[-.\s]?([0-9]{4})',  # Indian format
            r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',  # US format
            r'(?:\+?[0-9]{1,3}[-.\s]?)?([0-9]{8,15})'  # Generic international
        ]
        phones = []
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    phone = '-'.join(filter(None, match))
                else:
                    phone = match
                if phone and len(phone.replace('-', '').replace(' ', '')) >= 8:
                    phones.append(phone)
        contact_info['phones'] = list(set(phones))
        
        # LinkedIn pattern (improved)
        linkedin_patterns = [
            r'linkedin\.com/in/([A-Za-z0-9-_.]+)',
            r'LinkedIn:\s*([A-Za-z0-9-_.]+)',
            r'linkedin\.com/profile/view\?id=([A-Za-z0-9-_.]+)'
        ]
        linkedin_profiles = []
        for pattern in linkedin_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            linkedin_profiles.extend(matches)
        contact_info['linkedin'] = list(set(linkedin_profiles))
        
        # GitHub pattern (improved)
        github_patterns = [
            r'github\.com/([A-Za-z0-9-_.]+)(?:/([A-Za-z0-9-_.]+))?',
            r'GitHub:\s*([A-Za-z0-9-_.]+)'
        ]
        github_profiles = []
        github_repos = []
        for pattern in github_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    if match[0]:
                        github_profiles.append(match[0])
                    if len(match) > 1 and match[1]:
                        github_repos.append(f"{match[0]}/{match[1]}")
                else:
                    github_profiles.append(match)
        
        contact_info['github_profiles'] = list(set(github_profiles))
        contact_info['github_repos'] = list(set(github_repos))
        
        # Add hyperlinks if provided
        if hyperlinks:
            contact_info['hyperlinks'] = {
                'github_links': [link['url'] for link in hyperlinks.get('github', [])],
                'linkedin_links': [link['url'] for link in hyperlinks.get('linkedin', [])],
                'research_publication_links': [link['url'] for link in hyperlinks.get('research_publications', [])],
                'certification_links': [link['url'] for link in hyperlinks.get('certifications', [])],
                'social_media_links': [link['url'] for link in hyperlinks.get('social_media', [])],
                'other_links': [link['url'] for link in hyperlinks.get('other', [])]
            }
        
        return contact_info
    
    def clean_text_for_splitting(self, text: str) -> str:
        """Clean text for better splitting while preserving important information."""
        # Preserve emails and URLs in a more structured way
        text = re.sub(r'https?://[^\s]+', lambda m: f' [URL:{m.group()}] ', text)
        text = re.sub(r'www\.[^\s]+', lambda m: f' [URL:{m.group()}] ', text)
        
        # Preserve emails
        text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 
                     lambda m: f' [EMAIL:{m.group()}] ', text)
        
        # Clean up extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def split_resume_into_sections(self, resume_text: str, source_file_path: Optional[str] = None) -> List[Document]:
        """Split resume into logical sections and return as Document objects."""
        documents = []
        
        # Extract hyperlinks if source is a PDF
        hyperlinks = {}
        if source_file_path and source_file_path.lower().endswith('.pdf'):
            hyperlinks = extract_hyperlinks_from_pdf(source_file_path)
        
        # Extract contact info with hyperlinks
        contact_info = self.extract_contact_info(resume_text, hyperlinks)
        
        # Create contact info document
        contact_content = self._format_contact_info(contact_info)
        contact_doc = Document(
            page_content=contact_content,
            metadata={
                'section': 'contact_info',
                'doc_type': 'resume',
                'source': source_file_path or 'unknown',
                **contact_info
            }
        )
        documents.append(contact_doc)
        
        # Improved section patterns
        sections = {
            'summary': [
                r'(?:SUMMARY|OBJECTIVE|PROFILE|ABOUT\s+ME)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'experience': [
                r'(?:WORK\s+EXPERIENCE|PROFESSIONAL\s+EXPERIENCE|EXPERIENCE)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'education': [
                r'EDUCATION[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'skills': [
                r'(?:TECHNICAL\s+SKILLS?|SKILLS?\s+SUMMARY|SKILLS?)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|PROJECTS|RESEARCH|PUBLICATIONS|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'projects': [
                r'PROJECTS?[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|RESEARCH|PUBLICATIONS|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'certifications': [
                r'(?:RELEVANT\s+COURSEWORK\s+AND\s+)?CERTIFICATIONS?[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'research': [
                r'(?:RESEARCH|RESEARCH\s+EXPERIENCE)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|PUBLICATIONS|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'publications': [
                r'PUBLICATIONS?[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|CERTIFICATIONS|ACHIEVEMENTS|SOCIAL))|$)',
            ],
            'social': [
                r'(?:SOCIAL\s+ENGAGEMENTS?|SOCIAL\s+MEDIA|SOCIAL\s+LINKS|SOCIAL)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|ACHIEVEMENTS))|$)',
            ],
            'achievements': [
                r'(?:ACHIEVEMENTS?|AWARDS?|HONORS?)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL))|$)',
            ],
            'workshops': [
                r'(?:WORKSHOPS?|TRAINING)[\s:]*(.+?)(?=\n\s*(?:[A-Z\s]{3,}(?:EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS))|$)',
            ]
        }
        
        # Clean text for better pattern matching
        cleaned_text = self.clean_text_for_splitting(resume_text)
        
        # Extract sections
        for section_name, patterns in sections.items():
            for pattern in patterns:
                matches = re.findall(pattern, resume_text, re.IGNORECASE | re.DOTALL)
                for match in matches:
                    section_text = match.strip()
                    if section_text and len(section_text) > 20:  # Avoid very short matches
                        # Extract contact info specific to this section
                        section_contact = self.extract_contact_info(section_text, hyperlinks)
                        
                        # Split section into chunks if it's too long
                        chunks = self.splitter.split_text(section_text)
                        for i, chunk in enumerate(chunks):
                            if chunk.strip():
                                doc = Document(
                                    page_content=chunk.strip(),
                                    metadata={
                                        'section': section_name,
                                        'doc_type': 'resume',
                                        'source': source_file_path or 'unknown',
                                        'chunk_index': i,
                                        'total_chunks': len(chunks),
                                        'has_github': len(section_contact['github_repos']) > 0,
                                        'has_linkedin': len(section_contact['linkedin']) > 0,
                                        'github_repos': section_contact['github_repos'],
                                        'section_emails': section_contact['emails'],
                                        'hyperlinks_in_section': section_contact.get('hyperlinks', {})
                                    }
                                )
                                documents.append(doc)
                        break  # Only use first match for each section
        
        logger.info(f"✓ Created {len(documents)} document chunks from resume sections")
        return documents
    
    def _format_contact_info(self, contact_info: Dict[str, Any]) -> str:
        """Format contact information as readable text."""
        formatted = []
        
        if contact_info.get('emails'):
            formatted.append(f"Emails: {', '.join(contact_info['emails'])}")
        
        if contact_info.get('phones'):
            formatted.append(f"Phones: {', '.join(contact_info['phones'])}")
        
        if contact_info.get('linkedin'):
            formatted.append(f"LinkedIn: {', '.join(contact_info['linkedin'])}")
        
        if contact_info.get('github_profiles'):
            formatted.append(f"GitHub Profiles: {', '.join(contact_info['github_profiles'])}")
        
        if contact_info.get('github_repos'):
            formatted.append(f"GitHub Repositories: {', '.join(contact_info['github_repos'])}")
        
        return "\n".join(formatted) if formatted else "No contact information extracted"


async def create_vector_embeddings(documents: List[Document]) -> Optional[FAISS]:
    """Create vector embeddings for the resume documents using Azure OpenAI."""
    if not embedding_model:
        logger.error("❌ Embedding model not available. Cannot create vector embeddings.")
        return None
    
    if not documents:
        logger.warning("⚠️  No documents provided for embedding creation.")
        return None
    
    try:
        logger.info(f"Creating embeddings for {len(documents)} documents...")
        
        # Create embeddings
        vectorstore = await FAISS.afrom_documents(documents, embedding_model)
        
        logger.info("✓ Vector embeddings created successfully")
        return vectorstore
        
    except Exception as e:
        logger.error(f"❌ Error creating vector embeddings: {e}")
        return None


async def process_resume_from_file(file_path: str) -> Dict[str, Any]:
    """Process resume from file with embeddings and return comprehensive results."""
    try:
        # Load documents
        docs = resume_loader(file_path)
        full_text = "\n".join(d.page_content for d in docs)
        
        if not full_text.strip():
            logger.warning("⚠️  No text content found in the resume file")
            return {'documents': [], 'vectorstore': None, 'contact_info': {}}
        
        # Process resume into sections
        processor = ResumeProcessor()
        split_docs = processor.split_resume_into_sections(full_text, source_file_path=file_path)
        
        # Create vector embeddings
        vectorstore = await create_vector_embeddings(split_docs)
        
        # Extract contact info
        contact_doc = next((d for d in split_docs if d.metadata.get('section') == 'contact_info'), None)
        contact_info = contact_doc.metadata if contact_doc else {}
        
        # Print summary
        print(f"\n{'='*50}")
        print(f"RESUME PROCESSING SUMMARY")
        print(f"{'='*50}")
        print(f"✓ File: {os.path.basename(file_path)}")
        print(f"✓ Pages loaded: {len(docs)}")
        print(f"✓ Document chunks created: {len(split_docs)}")
        print(f"✓ Vector embeddings: {'Created' if vectorstore else 'Failed'}")
        
        # Show contact info
        if contact_info:
            print(f"\n📧 CONTACT INFORMATION:")
            if contact_info.get('emails'):
                print(f"  Email(s): {', '.join(contact_info['emails'])}")
            if contact_info.get('phones'):
                print(f"  Phone(s): {', '.join(contact_info['phones'])}")
            if contact_info.get('linkedin'):
                print(f"  LinkedIn: {', '.join(contact_info['linkedin'])}")
            if contact_info.get('github_profiles'):
                print(f"  GitHub: {', '.join(contact_info['github_profiles'])}")
        
        # Show hyperlinks summary
        if contact_info.get('hyperlinks'):
            hyperlinks = contact_info['hyperlinks']
            total_links = sum(len(links) for links in hyperlinks.values())
            if total_links > 0:
                print(f"\n🔗 HYPERLINKS EXTRACTED: {total_links} total")
                for category, links in hyperlinks.items():
                    if links:
                        print(f"  {category.replace('_', ' ').title()}: {len(links)}")
        
        # Show sections summary
        sections = {}
        for d in split_docs[1:]:  # Skip contact_info
            section = d.metadata.get('section', 'unknown')
            sections[section] = sections.get(section, 0) + 1
        
        if sections:
            print(f"\n📝 SECTIONS FOUND:")
            for section, count in sections.items():
                print(f"  {section.replace('_', ' ').title()}: {count} chunks")
        
        print(f"{'='*50}")
        
        return {
            'documents': split_docs,
            'vectorstore': vectorstore,
            'contact_info': contact_info,
            'sections': sections,
            'hyperlinks': contact_info.get('hyperlinks', {}),
            'file_path': file_path
        }
        
    except Exception as e:
        logger.error(f"❌ Error processing resume: {e}")
        return {'documents': [], 'vectorstore': None, 'contact_info': {}, 'error': str(e)}


async def search_resume_content(vectorstore: FAISS, query: str, k: int = 3) -> List[Document]:
    """Search resume content using vector similarity."""
    if not vectorstore:
        logger.error("❌ No vectorstore available for searching")
        return []
    
    try:
        results = await vectorstore.asimilarity_search(query, k=k)
        logger.info(f"✓ Found {len(results)} relevant documents for query: '{query}'")
        return results
    except Exception as e:
        logger.error(f"❌ Error searching resume content: {e}")
        return []


def get_resume_file_from_user() -> Optional[str]:
    """Interactive function to get resume file path from user with better guidance."""
    print("\n" + "="*60)
    print("📄 RESUME PROCESSOR - FILE INPUT")
    print("="*60)
    print("Supported formats: PDF (.pdf), Text (.txt), Markdown (.md), Word (.docx)")
    print("For hyperlink extraction, PDF format is recommended.")
    print()
    
    # Show some example paths
    print("💡 Example paths:")
    print("  Windows: C:\\Users\\YourName\\Documents\\resume.pdf")
    print("  Linux/Mac: /home/username/documents/resume.pdf")
    print("  Relative: ./resume.pdf")
    print()
    
    while True:
        file_path = input("📁 Enter resume file path (or 'quit' to exit): ").strip()
        
        if file_path.lower() in ['quit', 'exit', 'q']:
            return None
            
        if not file_path:
            print("❌ Please enter a file path.")
            continue
            
        # Handle quoted paths and clean input
        file_path = file_path.strip('"\'').strip()
        
        # Handle PowerShell execution artifacts
        if file_path.startswith('& '):
            file_path = file_path[2:].strip().strip('"\'')
        
        # Convert to absolute path if needed
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
            
        if os.path.exists(file_path):
            print(f"✓ File found: {file_path}")
            return file_path
        else:
            print(f"❌ File not found: {file_path}")
            print("   Please check the path and try again.")
            
            # Suggest similar files if directory exists
            dir_path = os.path.dirname(file_path)
            if os.path.exists(dir_path):
                pdf_files = [f for f in os.listdir(dir_path) if f.lower().endswith(('.pdf', '.txt', '.docx', '.md'))]
                if pdf_files:
                    print(f"   📂 Found these resume files in {dir_path}:")
                    for pdf_file in pdf_files[:5]:  # Show max 5 files
                        print(f"      - {pdf_file}")
            continue


async def demo_resume_search(result: Dict[str, Any]):
    """Demonstrate searching functionality with the processed resume."""
    if not result.get('vectorstore'):
        print("⚠️  Vector search not available (no embeddings created)")
        return
    
    print(f"\n🔍 DEMO: SEARCHING RESUME CONTENT")
    print(f"{'='*50}")
    
    # Example queries
    demo_queries = [
        "programming skills and technologies",
        "research experience and projects",
        "internship and work experience",
        "education and academic background"
    ]
    
    for query in demo_queries:
        print(f"\n🔎 Query: '{query}'")
        print("-" * 40)
        
        results = await search_resume_content(result['vectorstore'], query, k=2)
        for i, doc in enumerate(results, 1):
            section = doc.metadata.get('section', 'unknown')
            print(f"  {i}. Section: {section.replace('_', ' ').title()}")
            preview = doc.page_content[:150] + "..." if len(doc.page_content) > 150 else doc.page_content
            print(f"     Content: {preview}")
            print()


def validate_installation():
    """Check if all required packages are installed."""
    print("🔧 CHECKING INSTALLATION...")
    print("-" * 30)
    
    required_packages = {
        'PyMuPDF (fitz)': FITZ_AVAILABLE,
        'Azure OpenAI': embedding_model is not None,
        'LangChain': True,  # If we got this far, langchain is available
    }
    
    all_good = True
    for package, available in required_packages.items():
        status = "✓ Available" if available else "❌ Missing"
        print(f"{package}: {status}")
        if not available:
            all_good = False
    
    if not all_good:
        print("\n📦 INSTALLATION COMMANDS:")
        if not FITZ_AVAILABLE:
            print("pip install PyMuPDF")
        if embedding_model is None:
            print("Check your .env file for Azure OpenAI credentials")
    
    print()
    return all_good


if __name__ == "__main__":
    import argparse
    import sys
    
    # Validate installation first
    validate_installation()
    
    parser = argparse.ArgumentParser(
        description="Advanced Resume Processor with Hyperlink Extraction and Vector Embeddings",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python improved_resume.py resume.pdf                    # Process specific file
  python improved_resume.py -i                           # Interactive mode
  python improved_resume.py resume.pdf --demo            # Process and run search demo
  python improved_resume.py resume.pdf --github          # Process and analyze GitHub links
  python improved_resume.py resume.pdf --demo --github   # Full analysis with demos
        """
    )
    
    parser.add_argument("file", nargs="?", help="Path to resume file (.pdf/.txt/.docx/.md)")
    parser.add_argument("-i", "--interactive", action="store_true", 
                       help="Interactive mode - prompts for file path")
    parser.add_argument("--demo", action="store_true",
                       help="Run search demonstration after processing")
    parser.add_argument("--github", action="store_true",
                       help="Analyze GitHub links found in resume")
    parser.add_argument("-v", "--verbose", action="store_true",
                       help="Enable verbose logging")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    async def main():
        if args.interactive:
            # Interactive mode
            file_path = get_resume_file_from_user()
            if file_path is None:
                print("👋 Goodbye!")
                return
        elif args.file:
            # Command line file argument
            file_path = args.file
            if not os.path.exists(file_path):
                print(f"❌ File not found: {file_path}")
                return
        else:
            # No arguments - default to interactive mode
            print("💡 No file specified. Entering interactive mode...")
            file_path = get_resume_file_from_user()
            if file_path is None:
                print("👋 Goodbye!")
                return
        
        # Process the resume (with or without GitHub analysis)
        if args.github:
            # Process resume with GitHub analysis
            result = await process_resume_with_github_analysis(file_path)
        else:
            # Standard resume processing
            result = await process_resume_from_file(file_path)
        
        if result.get('error'):
            print(f"❌ Processing failed: {result['error']}")
            return
        
        # Run demo if requested
        if args.demo and result.get('vectorstore'):
            await demo_resume_search(result)
        
        # Display GitHub analysis summary if available
        if result.get('github_analysis') and result['github_analysis'].get('analysis'):
            github_summary = result['github_analysis']['summary']
            print(f"\n🔗 GITHUB ANALYSIS SUMMARY:")
            print(f"   Links found: {github_summary['total_links']}")
            print(f"   Profiles: {github_summary['profiles_found']}")
            print(f"   Repositories: {github_summary['repositories_found']}")
            print(f"   Unique users: {github_summary['unique_users']}")
        
        # Save vectorstore if created
        if result.get('vectorstore'):
            output_dir = os.path.dirname(file_path)
            vectorstore_path = os.path.join(output_dir, "resume_vectorstore")
            try:
                result['vectorstore'].save_local(vectorstore_path)
                print(f"💾 Vector store saved to: {vectorstore_path}")
            except Exception as e:
                logger.warning(f"⚠️  Could not save vectorstore: {e}")
        
        print(f"\n✅ Resume processing complete!")
        return result


# GitHub Link Parser Integration
class GitHubLinkParser:
    """Parse and analyze GitHub links extracted from resumes."""
    
    def __init__(self, github_token: Optional[str] = None):
        self.github_token = github_token or os.getenv("GITHUB_TOKEN")
        self.base_url = "https://api.github.com"
        self.headers = {
            "Accept": "application/vnd.github.v3+json",
            "User-Agent": "Resume-GitHub-Parser/1.0"
        }
        if self.github_token:
            self.headers["Authorization"] = f"token {self.github_token}"
    
    def extract_github_info_from_url(self, url: str) -> Dict[str, str]:
        """Extract GitHub username and repository from URL."""
        github_info = {'username': None, 'repository': None, 'url_type': 'unknown'}
        
        # Clean the URL
        url = url.strip().rstrip('/')
        
        # GitHub profile patterns
        profile_patterns = [
            r'github\.com/([A-Za-z0-9][-A-Za-z0-9]*[A-Za-z0-9])/?$',
            r'github\.com/([A-Za-z0-9][-A-Za-z0-9]*[A-Za-z0-9])/?\?.*$'
        ]
        
        # GitHub repository patterns
        repo_patterns = [
            r'github\.com/([A-Za-z0-9][-A-Za-z0-9]*[A-Za-z0-9])/([A-Za-z0-9][-A-Za-z0-9._]*[A-Za-z0-9])/?$',
            r'github\.com/([A-Za-z0-9][-A-Za-z0-9]*[A-Za-z0-9])/([A-Za-z0-9][-A-Za-z0-9._]*[A-Za-z0-9])/.*$'
        ]
        
        # Check for repository patterns first (more specific)
        for pattern in repo_patterns:
            match = re.search(pattern, url, re.IGNORECASE)
            if match:
                github_info['username'] = match.group(1)
                github_info['repository'] = match.group(2)
                github_info['url_type'] = 'repository'
                return github_info
        
        # Check for profile patterns
        for pattern in profile_patterns:
            match = re.search(pattern, url, re.IGNORECASE)
            if match:
                github_info['username'] = match.group(1)
                github_info['url_type'] = 'profile'
                return github_info
        
        return github_info
    
    async def fetch_github_profile(self, username: str) -> Dict[str, Any]:
        """Fetch GitHub profile information."""
        if not username:
            return {'error': 'No username provided'}
        
        url = f"{self.base_url}/users/{username}"
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, headers=self.headers)
                
                if response.status_code == 200:
                    data = response.json()
                    return {
                        'username': data.get('login'),
                        'name': data.get('name'),
                        'bio': data.get('bio'),
                        'company': data.get('company'),
                        'location': data.get('location'),
                        'email': data.get('email'),
                        'blog': data.get('blog'),
                        'twitter_username': data.get('twitter_username'),
                        'public_repos': data.get('public_repos', 0),
                        'followers': data.get('followers', 0),
                        'following': data.get('following', 0),
                        'created_at': data.get('created_at'),
                        'updated_at': data.get('updated_at'),
                        'profile_url': data.get('html_url')
                    }
                elif response.status_code == 404:
                    return {'error': f'GitHub user "{username}" not found'}
                else:
                    return {'error': f'GitHub API error: {response.status_code}'}
        except Exception as e:
            return {'error': f'Failed to fetch GitHub profile: {str(e)}'}
    
    async def fetch_github_repositories(self, username: str, limit: int = 10) -> List[Dict[str, Any]]:
        """Fetch user's repositories."""
        if not username:
            return []
        
        url = f"{self.base_url}/users/{username}/repos"
        params = {
            'sort': 'updated',
            'direction': 'desc',
            'per_page': limit
        }
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, headers=self.headers, params=params)
                
                if response.status_code == 200:
                    repos = response.json()
                    return [
                        {
                            'name': repo.get('name'),
                            'full_name': repo.get('full_name'),
                            'description': repo.get('description'),
                            'language': repo.get('language'),
                            'stars': repo.get('stargazers_count', 0),
                            'forks': repo.get('forks_count', 0),
                            'issues': repo.get('open_issues_count', 0),
                            'created_at': repo.get('created_at'),
                            'updated_at': repo.get('updated_at'),
                            'url': repo.get('html_url'),
                            'clone_url': repo.get('clone_url'),
                            'topics': repo.get('topics', []),
                            'is_private': repo.get('private', False),
                            'is_fork': repo.get('fork', False)
                        }
                        for repo in repos
                    ]
                else:
                    logger.warning(f"Failed to fetch repositories for {username}: {response.status_code}")
                    return []
        except Exception as e:
            logger.error(f"Error fetching repositories: {str(e)}")
            return []
    
    async def fetch_repository_details(self, username: str, repo_name: str) -> Dict[str, Any]:
        """Fetch detailed information about a specific repository."""
        if not username or not repo_name:
            return {'error': 'Username and repository name required'}
        
        url = f"{self.base_url}/repos/{username}/{repo_name}"
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, headers=self.headers)
                
                if response.status_code == 200:
                    repo = response.json()
                    return {
                        'name': repo.get('name'),
                        'full_name': repo.get('full_name'),
                        'description': repo.get('description'),
                        'language': repo.get('language'),
                        'languages_url': repo.get('languages_url'),
                        'stars': repo.get('stargazers_count', 0),
                        'forks': repo.get('forks_count', 0),
                        'watchers': repo.get('watchers_count', 0),
                        'issues': repo.get('open_issues_count', 0),
                        'created_at': repo.get('created_at'),
                        'updated_at': repo.get('updated_at'),
                        'pushed_at': repo.get('pushed_at'),
                        'url': repo.get('html_url'),
                        'clone_url': repo.get('clone_url'),
                        'topics': repo.get('topics', []),
                        'license': repo.get('license', {}).get('name') if repo.get('license') else None,
                        'is_private': repo.get('private', False),
                        'is_fork': repo.get('fork', False),
                        'is_archived': repo.get('archived', False),
                        'default_branch': repo.get('default_branch'),
                        'size': repo.get('size', 0)
                    }
                elif response.status_code == 404:
                    return {'error': f'Repository "{username}/{repo_name}" not found'}
                else:
                    return {'error': f'GitHub API error: {response.status_code}'}
        except Exception as e:
            return {'error': f'Failed to fetch repository details: {str(e)}'}
    
    async def analyze_github_links(self, github_links: List[str]) -> Dict[str, Any]:
        """Analyze all GitHub links and return comprehensive information."""
        results = {
            'profiles': {},
            'repositories': {},
            'analysis_summary': {
                'total_links': len(github_links),
                'unique_users': set(),
                'repositories_found': 0,
                'profiles_found': 0,
                'errors': []
            }
        }
        
        for link in github_links:
            try:
                github_info = self.extract_github_info_from_url(link)
                username = github_info.get('username')
                repository = github_info.get('repository')
                url_type = github_info.get('url_type')
                
                if not username:
                    results['analysis_summary']['errors'].append(f"Could not extract username from: {link}")
                    continue
                
                results['analysis_summary']['unique_users'].add(username)
                
                if url_type == 'repository' and repository:
                    # Fetch repository details
                    repo_details = await self.fetch_repository_details(username, repository)
                    if 'error' not in repo_details:
                        results['repositories'][f"{username}/{repository}"] = repo_details
                        results['analysis_summary']['repositories_found'] += 1
                    else:
                        results['analysis_summary']['errors'].append(f"Repository error: {repo_details['error']}")
                
                elif url_type == 'profile':
                    # Fetch profile details if not already fetched
                    if username not in results['profiles']:
                        profile_details = await self.fetch_github_profile(username)
                        if 'error' not in profile_details:
                            results['profiles'][username] = profile_details
                            results['analysis_summary']['profiles_found'] += 1
                            
                            # Also fetch user's top repositories
                            repos = await self.fetch_github_repositories(username, limit=5)
                            results['profiles'][username]['top_repositories'] = repos
                        else:
                            results['analysis_summary']['errors'].append(f"Profile error: {profile_details['error']}")
                
            except Exception as e:
                results['analysis_summary']['errors'].append(f"Error processing {link}: {str(e)}")
        
        # Convert set to count
        results['analysis_summary']['unique_users'] = len(results['analysis_summary']['unique_users'])
        
        return results


async def parse_github_links_from_resume(resume_result: Dict[str, Any]) -> Dict[str, Any]:
    """Extract and analyze GitHub links from resume processing result."""
    github_parser = GitHubLinkParser()
    
    # Collect all GitHub links from the resume result
    github_links = []
    
    # From contact info
    contact_info = resume_result.get('contact_info', {})
    
    # Direct GitHub links from hyperlinks
    hyperlinks = contact_info.get('hyperlinks', {})
    if hyperlinks.get('github_links'):
        github_links.extend(hyperlinks['github_links'])
    
    # GitHub profiles from contact extraction
    if contact_info.get('github_profiles'):
        for profile in contact_info['github_profiles']:
            github_links.append(f"https://github.com/{profile}")
    
    # GitHub repositories from contact extraction
    if contact_info.get('github_repos'):
        for repo in contact_info['github_repos']:
            github_links.append(f"https://github.com/{repo}")
    
    # Remove duplicates while preserving order
    seen = set()
    unique_github_links = []
    for link in github_links:
        if link not in seen:
            seen.add(link)
            unique_github_links.append(link)
    
    if not unique_github_links:
        return {
            'message': 'No GitHub links found in resume',
            'github_links': [],
            'analysis': {}
        }
    
    print(f"\n🔍 ANALYZING GITHUB LINKS")
    print(f"{'='*50}")
    print(f"Found {len(unique_github_links)} GitHub links:")
    for i, link in enumerate(unique_github_links, 1):
        print(f"  {i}. {link}")
    
    # Analyze all GitHub links
    analysis = await github_parser.analyze_github_links(unique_github_links)
    
    # Display results
    print(f"\n📊 GITHUB ANALYSIS RESULTS")
    print(f"{'='*50}")
    print(f"✓ Unique users found: {analysis['analysis_summary']['unique_users']}")
    print(f"✓ Profiles analyzed: {analysis['analysis_summary']['profiles_found']}")
    print(f"✓ Repositories analyzed: {analysis['analysis_summary']['repositories_found']}")
    
    if analysis['analysis_summary']['errors']:
        print(f"⚠️  Errors encountered: {len(analysis['analysis_summary']['errors'])}")
        for error in analysis['analysis_summary']['errors'][:3]:  # Show first 3 errors
            print(f"   - {error}")
    
    # Display profile summaries
    if analysis['profiles']:
        print(f"\n👤 GITHUB PROFILES:")
        for username, profile in analysis['profiles'].items():
            print(f"\n  📋 {username}")
            if profile.get('name'):
                print(f"      Name: {profile['name']}")
            if profile.get('bio'):
                print(f"      Bio: {profile['bio'][:100]}...")
            if profile.get('company'):
                print(f"      Company: {profile['company']}")
            if profile.get('location'):
                print(f"      Location: {profile['location']}")
            print(f"      Public Repos: {profile.get('public_repos', 0)}")
            print(f"      Followers: {profile.get('followers', 0)}")
            
            # Show top repositories
            top_repos = profile.get('top_repositories', [])
            if top_repos:
                print(f"      Top Repositories:")
                for repo in top_repos[:3]:
                    stars = repo.get('stars', 0)
                    language = repo.get('language', 'N/A')
                    print(f"        - {repo['name']} ({language}) ⭐{stars}")
    
    # Display repository details
    if analysis['repositories']:
        print(f"\n📁 REPOSITORY DETAILS:")
        for repo_name, repo in analysis['repositories'].items():
            print(f"\n  📂 {repo_name}")
            if repo.get('description'):
                print(f"      Description: {repo['description'][:100]}...")
            print(f"      Language: {repo.get('language', 'N/A')}")
            print(f"      Stars: {repo.get('stars', 0)} | Forks: {repo.get('forks', 0)}")
            if repo.get('topics'):
                print(f"      Topics: {', '.join(repo['topics'][:5])}")
            print(f"      Last Updated: {repo.get('updated_at', 'N/A')}")
    
    return {
        'github_links': unique_github_links,
        'analysis': analysis,
        'summary': {
            'total_links': len(unique_github_links),
            'profiles_found': analysis['analysis_summary']['profiles_found'],
            'repositories_found': analysis['analysis_summary']['repositories_found'],
            'unique_users': analysis['analysis_summary']['unique_users']
        }
    }


async def process_resume_with_github_analysis(file_path: str) -> Dict[str, Any]:
    """Process resume and automatically analyze GitHub links found."""
    try:
        # First, process the resume
        print("📄 Processing resume...")
        resume_result = await process_resume_from_file(file_path)
        
        if resume_result.get('error'):
            return resume_result
        
        # Then analyze GitHub links
        print("\n🔗 Analyzing GitHub links...")
        github_analysis = await parse_github_links_from_resume(resume_result)
        
        # Combine results
        combined_result = {
            **resume_result,
            'github_analysis': github_analysis
        }
        
        return combined_result
        
    except Exception as e:
        logger.error(f"❌ Error in combined processing: {e}")
        return {'error': str(e)}
    
    # Run the async main function
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n⏹️  Process interrupted by user")
    except Exception as e:
        logger.error(f"❌ Unexpected error: {e}")
        sys.exit(1)