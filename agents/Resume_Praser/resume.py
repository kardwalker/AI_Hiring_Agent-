from typing import List, Optional, Dict
import os
import asyncio
import logging
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader
from langchain_openai import AzureChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document
import re

# Optional PyMuPDF import for hyperlink extraction
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False
    logging.warning("PyMuPDF (fitz) not available. Install with 'pip install PyMuPDF' for hyperlink extraction.")
# Load environment variables once
load_dotenv()

AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY") or os.getenv("AZURE_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT") or os.getenv("AZURE_ENDPOINT")


if not AZURE_OPENAI_API_KEY or not AZURE_OPENAI_ENDPOINT:
    logging.warning("Azure OpenAI credentials not fully set. Set AZURE_OPENAI_API_KEY and AZURE_OPENAI_ENDPOINT in .env")

model = AzureChatOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    deployment_name="gpt-4o-mini",  # Specific deployment name in Azure
    api_version="2024-05-01-preview",
    temperature=0.2,
    max_tokens=512
)


def extract_hyperlinks_from_pdf(file_path: str) -> Dict[str, List[str]]:
    """Extract hyperlinks from PDF using PyMuPDF (fitz).
    
    Returns:
        Dict with categorized links: {
            'github': [...],
            'linkedin': [...],
            'research_publications': [...],
            'certifications': [...],
            'social_media': [...],
            'other': [...]
        }

"""
    if not FITZ_AVAILABLE:
        logging.warning("PyMuPDF not available. Cannot extract hyperlinks from PDF.")
        return {
            'github': [],
            'linkedin': [],
            'research_publications': [],
            'certifications': [],
            'social_media': [],
            'other': []
        }
    
    if not file_path.lower().endswith('.pdf'):
        return {
            'github': [],
            'linkedin': [],
            'research_publications': [],
            'certifications': [],
            'social_media': [],
            'other': []
        }
    
    links = {
        'github': [],
        'linkedin': [],
        'research_publications': [],
        'certifications': [],
        'social_media': [],
        'other': []
    }
    
    try:
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            link_list = page.get_links()
            
            for link in link_list:
                if 'uri' in link:
                    url = link['uri']
                    categorized_url = categorize_link(url)
                    category = categorized_url['category']
                    
                    if category in links:
                        links[category].append({
                            'url': url,
                            'page': page_num + 1,
                            'description': categorized_url['description'],
                            'rect': link.get('from', None)  # Position on page
                        })
        
        doc.close()
        
        # Remove duplicates while preserving order
        for category in links:
            seen = set()
            unique_links = []
            for link in links[category]:
                url = link['url']
                if url not in seen:
                    seen.add(url)
                    unique_links.append(link)
            links[category] = unique_links
        
        return links
        
    except Exception as e:
        logging.error(f"Error extracting hyperlinks from PDF: {e}")
        return {}


def categorize_link(url: str) -> Dict[str, str]:
    """Categorize a URL into relevant resume sections.
    
    Returns:
        Dict with 'category' and 'description'
    """
    url_lower = url.lower()
    
    # GitHub links
    if 'github.com' in url_lower:
        if '/repositories' in url_lower or '/repos' in url_lower:
            return {'category': 'github', 'description': 'GitHub repositories'}
        else:
            return {'category': 'github', 'description': 'GitHub profile/repository'}
    
    # LinkedIn links
    elif 'linkedin.com' in url_lower:
        return {'category': 'linkedin', 'description': 'LinkedIn profile'}
    
    # Research and publication platforms
    elif any(platform in url_lower for platform in [
        'researchgate.net', 'scholar.google.', 'orcid.org', 'ieee.org',
        'acm.org', 'arxiv.org', 'pubmed.ncbi.nlm.nih.gov', 'semanticscholar.org',
        'dblp.org', 'springer.com', 'sciencedirect.com', 'jstor.org'
    ]):
        return {'category': 'research_publications', 'description': 'Research/Publication platform'}
    
    # Certification platforms
    elif any(cert_platform in url_lower for cert_platform in [
        'credly.com', 'badgr.com', 'certifications.aws', 'cloud.google.com/certification',
        'docs.microsoft.com/certifications', 'coursera.org/account/accomplishments',
        'udacity.com/certificate', 'edx.org/certificates', 'codecademy.com/profiles',
        'freecodecamp.org/certification', 'cisco.com/c/en/us/training-events/training-certifications'
    ]):
        return {'category': 'certifications', 'description': 'Certification platform'}
    
    # Social media platforms
    elif any(social in url_lower for social in [
        'twitter.com', 'facebook.com', 'instagram.com', 'youtube.com',
        'medium.com', 'dev.to', 'stackoverflow.com', 'reddit.com'
    ]):
        return {'category': 'social_media', 'description': 'Social media profile'}
    
    # Portfolio/personal websites (heuristic)
    elif any(indicator in url_lower for indicator in [
        'portfolio', 'blog', 'personal', 'website'
    ]) or url_lower.count('.') == 1:  # Simple domain like johndoe.com
        return {'category': 'other', 'description': 'Personal website/portfolio'}
    
    else:
        return {'category': 'other', 'description': 'External link'}


def resume_loader(file_path: str) -> List[Document]:
    """Load resume from various file formats (.pdf, .txt, .docx)"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext in [".txt", ".md"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return [Document(page_content=content, metadata={"source": file_path})]
    elif ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return [Document(page_content=text, metadata={"source": file_path})]
        except ImportError:
            raise ImportError("python-docx package required for .docx files. Install with: pip install python-docx")
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX file: {e}")
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: .pdf, .txt, .md, .docx")

class ResumeProcessor:
    def __init__(self):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=50
        )
        
    def extract_contact_info(self, text, hyperlinks=None):
        """Extract and return structured contact information"""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['emails'] = emails
        
        # Phone pattern (various formats)
        phone_pattern = r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        contact_info['phones'] = ['-'.join(phone) for phone in phones]
        
        # LinkedIn pattern
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/profile/view\?id=)([A-Za-z0-9-_.]+)'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        contact_info['linkedin'] = linkedin
        
        # GitHub pattern
        github_pattern = r'(?:github\.com/)([A-Za-z0-9-_.]+)(?:/([A-Za-z0-9-_.]+))?'
        github = re.findall(github_pattern, text, re.IGNORECASE)
        contact_info['github_profiles'] = [match[0] for match in github]
        contact_info['github_repos'] = [f"{match[0]}/{match[1]}" for match in github if match[1]]
        
        # Add hyperlinks if provided
        if hyperlinks:
            contact_info['hyperlinks'] = {
                'github_links': [link['url'] for link in hyperlinks.get('github', [])],
                'linkedin_links': [link['url'] for link in hyperlinks.get('linkedin', [])],
                'research_publication_links': [link['url'] for link in hyperlinks.get('research_publications', [])],
                'certification_links': [link['url'] for link in hyperlinks.get('certifications', [])],
                'social_media_links': [link['url'] for link in hyperlinks.get('social_media', [])],
                'other_links': [link['url'] for link in hyperlinks.get('other', [])]
            }
        
        return contact_info
    
    def clean_text_for_splitting(self, text):
        """Remove or replace URLs/emails with placeholders to avoid splitting issues"""
        # Replace URLs with placeholders
        text = re.sub(r'https?://[^\s]+', '[URL]', text)
        text = re.sub(r'www\.[^\s]+', '[URL]', text)
        
        # Keep emails but ensure they don't break splitting
        text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', 
                     lambda m: f"[EMAIL: {m.group()}]", text)
        
        return text
    
    def split_resume(self, resume_text, source_file_path=None):
        documents = []
        
        # Extract hyperlinks if source is a PDF
        hyperlinks = {}
        if source_file_path and source_file_path.lower().endswith('.pdf'):
            hyperlinks = extract_hyperlinks_from_pdf(source_file_path)
        
        # First extract contact info with hyperlinks
        contact_info = self.extract_contact_info(resume_text, hyperlinks)
        
        # Create a contact info document
        contact_doc = Document(
            page_content=str(contact_info),
            metadata={
                'section': 'contact_info',
                'doc_type': 'resume',
                **contact_info
            }
        )
        documents.append(contact_doc)
        
        # Define sections with better patterns
        sections = {
            'summary': r'(?:SUMMARY|OBJECTIVE|PROFILE).*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?))',
            'experience': r'(?:WORK\s+|PROFESSIONAL\s+)?EXPERIENCE.*?(?=\n(?:[A-Z\s]{3,}|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'education': r'EDUCATION.*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'skills': r'(?:TECHNICAL\s+)?SKILLS?.*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'projects': r'PROJECTS?.*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'certifications': r'CERTIFICATIONS?.*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'research': r'(?:RESEARCH|RESEARCH\s+EXPERIENCE).*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|PUBLICATIONS|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'publications': r'PUBLICATIONS?.*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|SOCIAL|ACHIEVEMENTS?)|\Z)',
            'social': r'(?:SOCIAL\s+MEDIA|SOCIAL\s+LINKS|SOCIAL).*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|ACHIEVEMENTS?)|\Z)',
            'achievements': r'(?:ACHIEVEMENTS?|AWARDS?|HONORS?).*?(?=\n(?:[A-Z\s]{3,}|EXPERIENCE|EDUCATION|SKILLS|PROJECTS|RESEARCH|PUBLICATIONS|SOCIAL)|\Z)'
        }
        
        # Clean text for better splitting
        cleaned_text = self.clean_text_for_splitting(resume_text)
        
        for section_name, pattern in sections.items():
            matches = re.findall(pattern, cleaned_text, re.IGNORECASE | re.DOTALL)
            for match in matches:
                section_text = match.strip()
                if section_text:
                    # Extract URLs and repos from this section
                    section_contact = self.extract_contact_info(section_text, hyperlinks)
                    
                    chunks = self.splitter.split_text(section_text)
                    for chunk in chunks:
                        doc = Document(
                            page_content=chunk,
                            metadata={
                                'section': section_name,
                                'doc_type': 'resume',
                                'has_github': len(section_contact['github_repos']) > 0,
                                'has_linkedin': len(section_contact['linkedin']) > 0,
                                'github_repos': section_contact['github_repos'],
                                'section_emails': section_contact['emails'],
                                'hyperlinks_in_section': section_contact.get('hyperlinks', {})
                            }
                        )
                        documents.append(doc)
        
        return documents


async def process_resume_from_file(file_path: str):
    """Process resume from file path with validation and error handling"""
    try:
        print(f"Loading resume from: {file_path}")
        docs = resume_loader(file_path)
        full_text = "\n".join(d.page_content for d in docs)
        
        if not full_text.strip():
            print("Warning: No text content found in the resume file")
            return []
        
        processor = ResumeProcessor()
        split_docs = processor.split_resume(full_text, source_file_path=file_path)
        
        print(f"✓ Successfully loaded {len(docs)} page(s) -> {len(split_docs)} logical chunks")
        
        # Show hyperlinks if extracted
        contact_doc = next((d for d in split_docs if d.metadata.get('section') == 'contact_info'), None)
        if contact_doc and 'hyperlinks' in contact_doc.metadata:
            hyperlinks = contact_doc.metadata['hyperlinks']
            total_links = sum(len(links) for links in hyperlinks.values())
            if total_links > 0:
                print(f"✓ Extracted {total_links} hyperlinks from PDF")
                for category, links in hyperlinks.items():
                    if links:
                        print(f"  - {category.replace('_', ' ').title()}: {len(links)} links")
        
        # Show contact info if found
        if contact_doc:
            print(f"✓ Contact info extracted: {contact_doc.page_content}")
        
        # Show section summary
        sections = {}
        for d in split_docs[1:]:  # Skip contact_info
            section = d.metadata.get('section', 'unknown')
            sections[section] = sections.get(section, 0) + 1
        
        if sections:
            print(f"✓ Sections found: {', '.join(f'{k}({v})' for k, v in sections.items())}")
        
        return split_docs
        
    except FileNotFoundError:
        print(f"❌ Error: Resume file not found: {file_path}")
        return []
    except ValueError as e:
        print(f"❌ Error: {e}")
        return []
    except Exception as e:
        print(f"❌ Unexpected error processing resume: {e}")
        return []


def get_resume_file_from_user():
    """Interactive function to get resume file path from user"""
    print("=== Resume Input ===")
    print("Please provide your resume file path.")
    print("Supported formats: PDF (.pdf), Text (.txt), Markdown (.md), Word (.docx)")
    print()
    
    while True:
        file_path = input("Enter resume file path (or 'quit' to exit): ").strip()
        
        if file_path.lower() == 'quit':
            return None
            
        if not file_path:
            print("Please enter a file path.")
            continue
            
        # Handle quoted paths
        file_path = file_path.strip('"\'')
        
        # Convert to absolute path if needed
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
            
        if os.path.exists(file_path):
            return file_path
        else:
            print(f"File not found: {file_path}")
            print("Please check the path and try again.")
            continue


if __name__ == "__main__":  # pragma: no cover
    import argparse, sys
    
    parser = argparse.ArgumentParser(description="Resume processor with multiple input options")
    parser.add_argument("file", nargs="?", help="Path to resume file (.pdf/.txt/.docx/.md)")
    parser.add_argument("-i", "--interactive", action="store_true", 
                       help="Interactive mode - prompts for file path")
    
    args = parser.parse_args()
    
    if args.interactive:
        # Interactive mode
        file_path = get_resume_file_from_user()
        if file_path is None:
            print("Exiting...")
            sys.exit(0)
    elif args.file:
        # Command line file argument
        file_path = args.file
    else:
        # No arguments - default to interactive mode
        print("No file specified. Entering interactive mode...")
        file_path = get_resume_file_from_user()
        if file_path is None:
            print("Exiting...")
            sys.exit(0)
    
    # Process the resume
    asyncio.run(process_resume_from_file(file_path))


